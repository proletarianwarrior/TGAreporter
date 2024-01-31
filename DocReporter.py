# -*- coding: utf-8 -*-
# @Time : 2024/1/31 21:35
# @Author : DanYang
# @File : DocReporter.py
# @Software : PyCharm
from pathlib import Path
import json

from docx import Document
from docx.shared import Inches

from Logger import setup_logger
from DataLoader import cfg
from DataPloter import DataPloter

logger = setup_logger('report_data')

doc = Document()


def create_doc(file_path, **kwargs):
    if not isinstance(file_path, Path):
        file_path = Path(file_path)

    ploter = DataPloter(file_path)
    ploter.plot_TG(**kwargs)
    ploter.plot_DTG(**kwargs)
    ploter.plot_TG_DTG(**kwargs)
    ploter.plot_unite(**kwargs)

    name = file_path.name.split('.')[0]
    with open(f'unite/{name}_unite.json', 'r') as file:
        data = json.load(file)[0]

    width = cfg.getfloat('DataPlot', 'width')
    height = cfg.getfloat('DataPlot', 'height')
    doc.add_picture('image/TG-T.jpg', width=Inches(6.5), height=Inches(height * 6.5 / width))

    if bool(data['TG']):
        table = doc.add_table(rows=len(data['TG']), cols=2)
        for row, (key, value) in zip(table.rows, data['TG'].items()):
            row.cells[0].text = key
            row.cells[1].text = f'{[round(i, 2) for i in value]}'

    doc.add_picture('image/DTG-T.jpg', width=Inches(6.5), height=Inches(height * 6.5 / width))
    table = doc.add_table(rows=len(data['DTG']), cols=2)
    for row, (key, value) in zip(table.rows, data['DTG'].items()):
        row.cells[0].text = key
        value = [str(round(i, 2)) for i in value]
        row.cells[1].text = f'{"~".join(value)}'

    doc.add_picture('image/TG-DTG-T.jpg', width=Inches(6.5), height=Inches(height * 6.5 / width))
    table = doc.add_table(rows=len(data['TG_DTG']), cols=2)
    for row, (key, value) in zip(table.rows, data['TG_DTG'].items()):
        row.cells[0].text = key
        if abs(value) < 1:
            text = '.3f'
        elif abs(value) >= 1:
            text = '.2f'
        if abs(value) < 0.01:
            text = '.3e'
        row.cells[1].text = f'{value: {text}}'

    doc.add_picture('image/coal_pyrolysis.jpg', width=Inches(6.5), height=Inches(height * 6.5 / width))
    del data['unite']['coal_pyrolysis']['x']
    del data['unite']['coal_pyrolysis']['y']
    table = doc.add_table(rows=len(data['unite']['coal_pyrolysis']), cols=2)
    for row, (key, value) in zip(table.rows, data['unite']['coal_pyrolysis'].items()):
        row.cells[0].text = key
        if abs(value) < 1:
            text = '.3f'
        elif abs(value) >= 1:
            text = '.2f'
        if abs(value) < 0.01:
            text = '.3e'
        row.cells[1].text = f'{value: {text}}'

    doc.add_picture('image/carbon_combustion.jpg', width=Inches(6.5), height=Inches(height * 6.5 / width))
    del data['unite']['carbon_combustion']['x']
    del data['unite']['carbon_combustion']['y']
    table = doc.add_table(rows=len(data['unite']['carbon_combustion']), cols=2)
    for row, (key, value) in zip(table.rows, data['unite']['carbon_combustion'].items()):
        row.cells[0].text = key
        if abs(value) < 1:
            text = '.3f'
        elif abs(value) >= 1:
            text = '.2f'
        if abs(value) < 0.01:
            text = '.3e'
        row.cells[1].text = f'{value: {text}}'

    doc.save('report.docx')
    logger.info(f'输出报告请见report.docx')


if __name__ == '__main__':
    file_path = './data/ymt.txt'
    create_doc(file_path)

